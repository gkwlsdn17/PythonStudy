import docx
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH

class Certification:

    def __init__(self, name, birth, ho):
        self.name = name
        self.birth = birth
        self.ho = ho

    def make(self):
        doc = docx.Document(r'template.docx')

        style = doc.styles['Normal']
        style.font.name = '나눔고딕'
        style._element.rPr.rFonts.set(qn('w:eastAsia'), '나눔고딕')
        style.font.size = docx.shared.Pt(12)

        para = doc.add_paragraph()
        run = para.add_run('\n\n')
        run = para.add_run(f' 제 {self.ho} 호\n')
        run.font.name = '나눔고딕'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '나눔고딕')
        run.font.size = docx.shared.Pt(20)

        para = doc.add_paragraph()
        run = para.add_run('\n')
        run = para.add_run('수 료 증')
        run.font.name = '나눔고딕'
        run.bold = True
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '나눔고딕')
        run.font.size = docx.shared.Pt(40)
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        para = doc.add_paragraph()
        run = para.add_run('\n\n')
        run = para.add_run(f' 성 명: {self.name}\n')
        run.font.name = '나눔고딕'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '나눔고딕')
        run.font.size = docx.shared.Pt(20)
        run = para.add_run(f' 생 년 월 일: {self.birth}\n')
        run.font.name = '나눔고딕'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '나눔고딕')
        run.font.size = docx.shared.Pt(20)
        run = para.add_run(f' 교 육 과 정: 파이썬과 40개의 작품들\n')
        run.font.name = '나눔고딕'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '나눔고딕')
        run.font.size = docx.shared.Pt(20)
        run = para.add_run(f' 교 육 날 짜: 2022-09-29\n')
        run.font.name = '나눔고딕'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '나눔고딕')
        run.font.size = docx.shared.Pt(20)

        para = doc.add_paragraph()
        run = para.add_run('\n\n')
        run = para.add_run('위 사람은 파이썬과 40개의 작품들 교육과정을\n')
        run.font.name = '나눔고딕'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '나눔고딕')
        run.font.size = docx.shared.Pt(20)
        run = para.add_run('이수하였으므로 이 증서를 수여 합니다.\n')
        run.font.name = '나눔고딕'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '나눔고딕')
        run.font.size = docx.shared.Pt(20)
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        para = doc.add_paragraph()
        run = para.add_run('\n\n')
        run = para.add_run('2022.09.29\n')
        run.font.name = '나눔고딕'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '나눔고딕')
        run.font.size = docx.shared.Pt(20)
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        para = doc.add_paragraph()
        run = para.add_run('\n\n\n')
        run = para.add_run('파이썬교육기관장')
        run.font.name = '나눔고딕'
        run.bold = True
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '나눔고딕')
        run.font.size = docx.shared.Pt(20)
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.save(f'{self.name} certification.docx')